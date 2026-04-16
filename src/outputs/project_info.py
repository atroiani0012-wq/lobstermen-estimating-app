"""
Output 2 — Project Information Word doc.

Sections:
  1. Project Summary
  2. File Appendix
  3. Testing Requirements
  4. Design Requirements (preliminary-design-ready)
  5. Unknowns / Assumptions (with reasoning)
  6. Estimating Risks
  7. Equipment List
  8. Vendor List
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


def _set_cell(cell, text: str, bold: bool = False, shade: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text if text is not None else "")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    if bold:
        run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if shade:
        tc_pr = cell._tc.get_or_add_tcPr()
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        tc_pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade}"/>'))


def _add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(11)


def _add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text if text else "—")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    if bold:
        run.bold = True


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for c, h in enumerate(headers):
        _set_cell(t.rows[0].cells[c], h, bold=True, shade="1E3A8A")
        t.rows[0].cells[c].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        tr = t.add_row()
        for c, val in enumerate(row):
            _set_cell(tr.cells[c], str(val) if val is not None else "")


def build_project_info_docx(analysis: dict[str, Any]) -> bytes:
    doc = Document()

    # Default font
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    p = doc.add_paragraph()
    run = p.add_run("UMA Project Information Brief")
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    project = analysis.get("project") or {}
    _add_heading(doc, "1. Project Summary", 1)
    _add_table(
        doc,
        ["Field", "Value"],
        [
            ["Project", project.get("name", "")],
            ["Location", project.get("location", "")],
            ["Client / GC", project.get("client", "")],
            ["Bid Due Date", project.get("bid_due_date", "")],
            ["Scope Type", project.get("scope_type", "")],
            ["Site Conditions", project.get("site_conditions", "")],
            ["Schedule Notes", project.get("schedule_notes", "")],
        ],
    )
    _add_heading(doc, "Description", 2)
    _add_para(doc, project.get("description", ""))

    _add_heading(doc, "2. File Appendix", 1)
    appendix = analysis.get("file_appendix") or []
    if appendix:
        _add_table(doc, ["File", "Type", "Summary", "Key info"],
                   [[f.get("filename", ""), f.get("type", ""), f.get("summary", ""), f.get("key_info", "")] for f in appendix])
    else:
        _add_para(doc, "No files were provided.")

    _add_heading(doc, "3. Testing Requirements", 1)
    testing = analysis.get("testing_requirements") or []
    if testing:
        _add_table(doc, ["Test", "Unit", "Qty", "Spec Section", "Notes"],
                   [[t.get("test_type", ""), t.get("unit", ""), t.get("quantity", 0),
                     t.get("reference_spec_section", ""), t.get("notes", "")] for t in testing])
    else:
        _add_para(doc, "None identified in the bid package.")

    _add_heading(doc, "4. Design Requirements", 1)
    design = analysis.get("design_requirements") or {}
    _add_table(doc, ["Field", "Value"],
               [["Design-Build", design.get("design_build", "")],
                ["Ready for prelim pricing?", design.get("design_ready_for_prelim_pricing", "")],
                ["Design responsibility", design.get("design_responsibility", "")]])
    _add_heading(doc, "Key design inputs needed", 2)
    for ki in design.get("key_design_inputs_needed") or []:
        _add_para(doc, f"• {ki}")
    _add_heading(doc, "Preliminary design notes (UMA assumptions for pricing)", 2)
    _add_para(doc, design.get("preliminary_design_notes", ""))

    _add_heading(doc, "5. Unknowns & Assumptions", 1)
    ua = analysis.get("unknowns_and_assumptions") or []
    if ua:
        _add_table(doc, ["Unknown", "Assumption", "Reasoning", "Impact if wrong"],
                   [[u.get("item", ""), u.get("assumption", ""), u.get("reasoning", ""), u.get("impact_if_wrong", "")] for u in ua])
    else:
        _add_para(doc, "None flagged.")

    _add_heading(doc, "6. Estimating Risks", 1)
    risks = analysis.get("estimating_risks") or []
    if risks:
        _add_table(doc, ["Risk", "Severity", "Mitigation"],
                   [[r.get("risk", ""), r.get("severity", ""), r.get("mitigation", "")] for r in risks])
    else:
        _add_para(doc, "No material risks flagged.")

    _add_heading(doc, "7. Equipment List", 1)
    equip = analysis.get("equipment_list") or []
    if equip:
        _add_table(doc, ["Equipment", "Duration (days)", "Notes"],
                   [[e.get("equipment", ""), e.get("duration_days", 0), e.get("notes", "")] for e in equip])
    else:
        _add_para(doc, "To be determined during pricing.")

    _add_heading(doc, "8. Vendor List", 1)
    vendors = analysis.get("vendor_list") or []
    if vendors:
        _add_table(doc, ["Category", "Suggested Vendors", "# takeoff items", "# testing items"],
                   [[v.get("vendor_category", ""), ", ".join(v.get("suggested_vendors") or []),
                     len(v.get("takeoff_items_for_rfq") or []),
                     len(v.get("testing_items_for_rfq") or [])] for v in vendors])
    else:
        _add_para(doc, "No vendors identified.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
