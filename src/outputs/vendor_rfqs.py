"""
Output 3 — Vendor RFQ drafts (one per vendor-category, packaged in a single .docx).

Each RFQ section has:
- Header block with project info
- Takeoff table (items that vendor can quote)
- Testing items broken out (if applicable)
- Standard terms stub (lead time, delivery, pricing basis)
"""
from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


def _set_cell(cell, text: str, bold: bool = False, shade: str | None = None, color: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text if text is not None else "")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if shade:
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade}"/>'))


def _add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"


def _add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    if bold:
        run.bold = True


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for c, h in enumerate(headers):
        _set_cell(t.rows[0].cells[c], h, bold=True, shade="1E3A8A", color="FFFFFF")
    for row in rows:
        tr = t.add_row()
        for c, val in enumerate(row):
            _set_cell(tr.cells[c], str(val) if val is not None else "")


def _project_header_block(doc: Document, project: dict[str, Any]):
    _add_table(doc, ["Field", "Value"], [
        ["Project", project.get("name", "")],
        ["Location", project.get("location", "")],
        ["Client / GC", project.get("client", "")],
        ["Bid Due Date", project.get("bid_due_date", "")],
        ["Scope Type", project.get("scope_type", "")],
    ])


def _vendor_section(doc: Document, project: dict[str, Any], vendor: dict[str, Any], is_first: bool):
    if not is_first:
        doc.add_page_break()
    category = vendor.get("vendor_category", "(unspecified category)")
    suggested = ", ".join(vendor.get("suggested_vendors") or []) or "(to be selected)"

    p = doc.add_paragraph()
    r = p.add_run(f"REQUEST FOR QUOTATION — {category}")
    r.font.name = "Arial"; r.font.size = Pt(16); r.bold = True

    _add_para(doc, f"Suggested vendors: {suggested}")
    _add_para(doc, f"Issued: {date.today().isoformat()}    Please respond by: (bid due date above)")
    _add_heading(doc, "Project", 2)
    _project_header_block(doc, project)

    _add_heading(doc, "Scope description", 2)
    _add_para(doc, project.get("description", ""))

    takeoff_items = vendor.get("takeoff_items_for_rfq") or []
    _add_heading(doc, "Items to quote", 2)
    if takeoff_items:
        _add_table(
            doc,
            ["Item / Material", "Unit", "Quantity", "Spec / notes"],
            [[i.get("item", ""), i.get("unit", ""), i.get("quantity", 0), i.get("spec_notes", "")] for i in takeoff_items],
        )
    else:
        _add_para(doc, "(no items assigned — confirm scope with estimator)")

    testing_items = vendor.get("testing_items_for_rfq") or []
    if testing_items:
        _add_heading(doc, "Testing items (please price separately)", 2)
        _add_table(doc, ["Test / Service"], [[t] for t in testing_items])

    _add_heading(doc, "Standard UMA RFQ terms", 2)
    for line in [
        "• Please quote unit prices in the table above (FOB jobsite unless noted).",
        "• Provide lead time for each item and any minimum-order constraints.",
        "• Identify any items that require submittal / mill cert / COA.",
        "• Note any exclusions, alternates, or value-engineering suggestions.",
        "• Quote valid for 30 days minimum.",
        "• Direct questions to the UMA estimator (contact info in email).",
    ]:
        _add_para(doc, line)


def build_vendor_rfqs_docx(analysis: dict[str, Any]) -> bytes:
    doc = Document()
    # Default font
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Cover
    p = doc.add_paragraph()
    r = p.add_run("Vendor RFQ Drafts")
    r.font.name = "Arial"; r.font.size = Pt(20); r.bold = True

    project = analysis.get("project") or {}
    _add_para(doc, f"Project: {project.get('name', '')}", bold=True)
    _add_para(doc, f"Location: {project.get('location', '')}")
    _add_para(doc, f"Bid due: {project.get('bid_due_date', '')}")

    vendors = analysis.get("vendor_list") or []
    if not vendors:
        _add_para(doc, "\nNo vendor categories identified in the analysis. Add vendors manually.")
    else:
        for i, v in enumerate(vendors):
            _vendor_section(doc, project, v, is_first=(i == 0))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
