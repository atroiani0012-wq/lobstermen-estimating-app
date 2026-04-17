"""File ingest: convert uploaded bid-package files into Claude-ready content blocks.

Handles PDF, DOCX, XLSX/XLS/CSV, images (PNG/JPG/GIF/WEBP), TXT, with a
fall-through block for unknown types. Mirrors the Streamlit version's
behavior so Claude sees the same shape of input.
"""
from __future__ import annotations

import base64
import io
import mimetypes
import pathlib
from dataclasses import dataclass, field
from typing import Any

MAX_PDF_DIRECT_BYTES = 30 * 1024 * 1024  # 30 MB — send natively as document block
MAX_PDF_PAGES_AS_IMAGES = 10
MAX_IMG_BYTES = 5 * 1024 * 1024          # 5 MB per image block


@dataclass
class IngestedFile:
    filename: str
    content_blocks: list[dict[str, Any]] = field(default_factory=list)
    text_summary: str = ""
    error: str | None = None


def _b64(data: bytes) -> str:
    return base64.standard_b64encode(data).decode("ascii")


def _image_block(data: bytes, media_type: str) -> dict[str, Any]:
    return {
        "type": "image",
        "source": {"type": "base64", "media_type": media_type, "data": _b64(data)},
    }


def _document_block(data: bytes, media_type: str = "application/pdf") -> dict[str, Any]:
    return {
        "type": "document",
        "source": {"type": "base64", "media_type": media_type, "data": _b64(data)},
    }


def _ingest_pdf(filename: str, data: bytes) -> IngestedFile:
    out = IngestedFile(filename=filename)
    size = len(data)
    if size <= MAX_PDF_DIRECT_BYTES:
        out.content_blocks.append(_document_block(data, "application/pdf"))
        out.text_summary = f"[PDF {filename}] attached as document block ({size/1024/1024:.1f} MB)"
        return out
    try:
        import pdfplumber
        text_chunks: list[str] = []
        page_images: list[bytes] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                txt = page.extract_text() or ""
                if txt.strip():
                    text_chunks.append(f"--- Page {i+1} ---\n{txt}")
                if i < MAX_PDF_PAGES_AS_IMAGES:
                    try:
                        im = page.to_image(resolution=100).original
                        buf = io.BytesIO()
                        im.save(buf, format="PNG")
                        page_images.append(buf.getvalue())
                    except Exception:
                        pass
        joined = "\n\n".join(text_chunks)[:60000]
        out.content_blocks.append({
            "type": "text",
            "text": f"[PDF {filename} — large, text-extracted from {total_pages} pages]\n\n{joined}",
        })
        for img in page_images:
            out.content_blocks.append(_image_block(img, "image/png"))
        out.text_summary = (
            f"[PDF {filename}] large ({size/1024/1024:.1f} MB, {total_pages} pp) — "
            f"extracted text + first {len(page_images)} pages as images"
        )
    except Exception as e:
        out.error = f"PDF extract failed: {e}"
        out.text_summary = f"[PDF {filename}] ERROR: {e}"
    return out


def _ingest_docx(filename: str, data: bytes) -> IngestedFile:
    out = IngestedFile(filename=filename)
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(parts)[:60000]
        out.content_blocks.append({"type": "text", "text": f"[DOCX {filename}]\n\n{text}"})
        out.text_summary = f"[DOCX {filename}] {len(parts)} blocks extracted"
    except Exception as e:
        out.error = f"DOCX extract failed: {e}"
        out.text_summary = f"[DOCX {filename}] ERROR: {e}"
    return out


def _ingest_excel(filename: str, data: bytes) -> IngestedFile:
    out = IngestedFile(filename=filename)
    try:
        import pandas as pd
        buf = io.BytesIO(data)
        xls = pd.ExcelFile(buf)
        parts: list[str] = [f"[XLSX {filename} — {len(xls.sheet_names)} sheets]"]
        for sheet in xls.sheet_names[:8]:
            df = pd.read_excel(xls, sheet_name=sheet, nrows=200, header=None)
            parts.append(f"\n=== Sheet: {sheet} ({df.shape[0]} rows x {df.shape[1]} cols shown) ===")
            parts.append(df.fillna("").to_csv(index=False, header=False))
        text = "\n".join(parts)[:60000]
        out.content_blocks.append({"type": "text", "text": text})
        out.text_summary = f"[XLSX {filename}] {len(xls.sheet_names)} sheets"
    except Exception as e:
        out.error = f"XLSX extract failed: {e}"
        out.text_summary = f"[XLSX {filename}] ERROR: {e}"
    return out


def _ingest_csv(filename: str, data: bytes) -> IngestedFile:
    out = IngestedFile(filename=filename)
    try:
        text = data.decode("utf-8", errors="replace")[:60000]
        out.content_blocks.append({"type": "text", "text": f"[CSV {filename}]\n{text}"})
        out.text_summary = f"[CSV {filename}] {len(text)} chars"
    except Exception as e:
        out.error = f"CSV read failed: {e}"
        out.text_summary = f"[CSV {filename}] ERROR: {e}"
    return out


def _ingest_image(filename: str, data: bytes, media_type: str) -> IngestedFile:
    out = IngestedFile(filename=filename)
    if len(data) > MAX_IMG_BYTES:
        try:
            from PIL import Image
            im = Image.open(io.BytesIO(data))
            im.thumbnail((2000, 2000))
            buf = io.BytesIO()
            fmt = "JPEG" if media_type in ("image/jpeg", "image/jpg") else "PNG"
            im.convert("RGB" if fmt == "JPEG" else "RGBA").save(
                buf, format=fmt, quality=85 if fmt == "JPEG" else None
            )
            data = buf.getvalue()
            media_type = f"image/{fmt.lower()}"
        except Exception:
            pass
    out.content_blocks.append({"type": "text", "text": f"[IMG {filename}]"})
    out.content_blocks.append(_image_block(data, media_type))
    out.text_summary = f"[IMG {filename}] {len(data)/1024:.0f} KB"
    return out


def _ingest_text(filename: str, data: bytes) -> IngestedFile:
    out = IngestedFile(filename=filename)
    try:
        text = data.decode("utf-8", errors="replace")[:60000]
        out.content_blocks.append({"type": "text", "text": f"[TXT {filename}]\n{text}"})
        out.text_summary = f"[TXT {filename}] {len(text)} chars"
    except Exception as e:
        out.error = f"TXT read failed: {e}"
        out.text_summary = f"[TXT {filename}] ERROR: {e}"
    return out


def ingest_file(filename: str, data: bytes) -> IngestedFile:
    ext = pathlib.Path(filename).suffix.lower()
    mime, _ = mimetypes.guess_type(filename)
    mime = (mime or "").lower()

    if ext == ".pdf" or mime == "application/pdf":
        return _ingest_pdf(filename, data)
    if ext == ".docx":
        return _ingest_docx(filename, data)
    if ext in (".xlsx", ".xlsm", ".xls"):
        return _ingest_excel(filename, data)
    if ext == ".csv":
        return _ingest_csv(filename, data)
    if ext == ".png" or mime == "image/png":
        return _ingest_image(filename, data, "image/png")
    if ext in (".jpg", ".jpeg") or mime in ("image/jpeg", "image/jpg"):
        return _ingest_image(filename, data, "image/jpeg")
    if ext == ".gif" or mime == "image/gif":
        return _ingest_image(filename, data, "image/gif")
    if ext == ".webp" or mime == "image/webp":
        return _ingest_image(filename, data, "image/webp")
    if ext in (".txt", ".md", ".log"):
        return _ingest_text(filename, data)

    out = IngestedFile(filename=filename)
    out.content_blocks.append({
        "type": "text",
        "text": (
            f"[UNSUPPORTED {filename} — {len(data)/1024:.0f} KB, ext={ext!r}] "
            "Skipped content extraction; ask the user about this file in your unknowns list."
        ),
    })
    out.text_summary = f"[UNSUPPORTED {filename}]"
    return out


def ingest_many(files: list[tuple[str, bytes]]) -> list[IngestedFile]:
    return [ingest_file(name, data) for name, data in files]
