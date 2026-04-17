"""Generate 04_Cost_Proposal.docx — first-draft cost proposal."""
from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor, Inches


def _set_cell(cell, text: str, bold: bool = False, shade: str | None = None, color: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text if text is not None else "")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if shade:
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade}"/>'))


def _add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"


def _add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text if text else "")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    if bold:
        run.bold = True


def _add_bullets(doc: Document, items: list[str]):
    for it in items or []:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.name = "Arial"
        run.font.size = Pt(11)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for c, h in enumerate(headers):
        _set_cell(t.rows[0].cells[c], h, bold=True, shade="1E3A8A", color="FFFFFF")
    for row in rows:
        tr = t.add_row()
        for c, val in enumerate(row):
            _set_cell(tr.cells[c], str(val) if val is not None else "")


def build_proposal_docx(analysis: dict[str, Any]) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    project = analysis.get("project") or {}
    cp = analysis.get("cost_proposal_draft") or {}

    p = doc.add_paragraph()
    r = p.add_run("COST PROPOSAL (First Draft)")
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.bold = True

    _add_para(doc, "Prepared by: UMA")
    _add_para(doc, f"Date: {date.today().isoformat()}")
    _add_para(doc, f"For: {project.get('client', '')}")
    _add_para(doc, f"Re: {project.get('name', '')} — {project.get('location', '')}")

    _add_heading(doc, "Executive Summary", 1)
    _add_para(doc, cp.get("executive_summary", ""))

    _add_heading(doc, "Scope of Work", 1)
    _add_para(doc, project.get("description", ""))
    _add_para(doc, f"Scope type: {project.get('scope_type', '')}")

    _add_heading(doc, "Inclusions", 1)
    _add_bullets(doc, cp.get("inclusions") or [])

    _add_heading(doc, "Exclusions", 1)
    _add_bullets(doc, cp.get("exclusions") or [])

    _add_heading(doc, "Clarifications", 1)
    _add_bullets(doc, cp.get("clarifications") or [])

    _add_heading(doc, "Commercial Terms", 1)
    _add_table(
        doc,
        ["Field", "Value"],
        [
            ["Pricing basis", cp.get("pricing_basis", "")],
            ["Payment terms", cp.get("payment_terms", "")],
            ["Contingency %", cp.get("contingency_pct", 0)],
            ["Markup %", cp.get("markup_pct", 0)],
            ["Bond required?", cp.get("bond_required", "")],
            ["Bid validity (days)", cp.get("bid_validity_days", 30)],
        ],
    )

    _add_heading(doc, "Pricing Placeholder", 1)
    _add_para(doc, "Line-item pricing to be finalized after review of the UMA Master Estimator output.")
    takeoff = analysis.get("takeoff_items") or []
    if takeoff:
        _add_table(
            doc,
            ["Item", "Unit", "Qty", "Unit Price ($)", "Extended ($)"],
            [
                [
                    t.get("item", ""),
                    t.get("unit", ""),
                    t.get("quantity", 0),
                    t.get("unit_cost_est", 0) or "TBD",
                    (t.get("quantity", 0) or 0) * (t.get("unit_cost_est", 0) or 0) or "TBD",
                ]
                for t in takeoff
            ],
        )

    _add_heading(doc, "Assumptions", 1)
    for u in analysis.get("unknowns_and_assumptions") or []:
        _add_para(doc, f"• {u.get('assumption', '')} — {u.get('reasoning', '')}")

    _add_heading(doc, "Testing (priced separately)", 1)
    testing = analysis.get("testing_requirements") or []
    if testing:
        _add_table(
            doc,
            ["Test", "Unit", "Qty", "Notes"],
            [
                [t.get("test_type", ""), t.get("unit", ""), t.get("quantity", 0), t.get("notes", "")]
                for t in testing
            ],
        )
    else:
        _add_para(doc, "None identified.")

    _add_heading(doc, "Sign-off", 1)
    _add_para(doc, "Accepted by: _______________________________    Date: ____________")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
